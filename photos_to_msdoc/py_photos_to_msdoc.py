import os
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

def create_word_document_with_images(folder_path, output_filename="images.docx"):
    """
    Creates a Word document with each image in the folder on its own page,
    preceded by the filename as a title, with proper image scaling.
    """
    # Supported image extensions
    image_extensions = ('.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp')
    
    # Get all image files in the folder
    image_files = [f for f in os.listdir(folder_path) 
                  if f.lower().endswith(image_extensions)]
    
    if not image_files:
        print("No image files found in the specified folder.")
        return
    
    # Create a new Word document
    doc = Document()
    
    # Set page margins (left, right, top, bottom) in inches
    section = doc.sections[0]
    section.left_margin = Cm(1.5)  # 1.5 cm margins
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)
    
    for image_file in image_files:
        # Add filename as title (without extension)
        title = os.path.splitext(image_file)[0]
        
        # Add title paragraph
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run(title)
        run.bold = True
        run.font.size = Pt(24)  # Slightly smaller font size
        
        # Add image
        img_path = os.path.join(folder_path, image_file)
        
        try:
            # Get image dimensions to maintain aspect ratio
            with Image.open(img_path) as img:
                img_width, img_height = img.size
                aspect_ratio = img_width / img_height
            
            # Calculate maximum dimensions that fit within page
            max_width = 6.0  # Inches (accounts for margins)
            max_height = 8.0  # Inches (accounts for header/footer space)
            
            # Calculate dimensions maintaining aspect ratio
            if aspect_ratio > (max_width/max_height):
                width = Inches(min(max_width, img_width/72))  # 72 dpi approximation
                height = width / aspect_ratio
            else:
                height = Inches(min(max_height, img_height/72))
                width = height * aspect_ratio
            
            # Add image paragraph
            img_para = doc.add_paragraph()
            img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Add image with calculated dimensions
            run = img_para.add_run()
            run.add_picture(img_path, width=width, height=height)
            
            # Add page break after image (unless it's the last one)
            if image_file != image_files[-1]:
                doc.add_page_break()
                
        except Exception as e:
            print(f"Could not add image {image_file}: {str(e)}")
            continue
    
    # Save the document
    doc.save(output_filename)
    print(f"Document created successfully: {output_filename}")

if __name__ == "__main__":
    folder_path = input("Enter the path to the folder containing images: ")
    output_filename = input("Enter the output filename (default: images.docx): ") or "images.docx"
    
    create_word_document_with_images(folder_path, output_filename)
